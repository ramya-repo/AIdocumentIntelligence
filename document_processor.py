import streamlit as st
import pandas as pd
import PyPDF2
import docx
from io import BytesIO
import os

class DocumentProcessor:
    def __init__(self):
        pass
    
    def process_file(self, uploaded_file):
        """Process uploaded file based on its type and return extracted content"""
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._process_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._process_word(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                return self._process_excel(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
            return None
    
    def _process_pdf(self, uploaded_file):
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text.strip()
        
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _process_word(self, uploaded_file):
        """Extract text from Word document"""
        try:
            doc = docx.Document(BytesIO(uploaded_file.read()))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from Word document")
            
            return text.strip()
        
        except Exception as e:
            raise Exception(f"Word document processing error: {str(e)}")
    
    def _process_excel(self, uploaded_file):
        """Extract data from Excel file and convert to readable text format"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(BytesIO(uploaded_file.read()))
            all_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                if df.empty:
                    continue
                
                sheet_content = f"\n=== Sheet: {sheet_name} ===\n"
                
                # Add column headers
                sheet_content += "Columns: " + ", ".join(df.columns.astype(str)) + "\n\n"
                
                # Add data summary
                sheet_content += f"Number of rows: {len(df)}\n"
                sheet_content += f"Number of columns: {len(df.columns)}\n\n"
                
                # Convert DataFrame to readable text format
                # For large datasets, limit the number of rows displayed
                max_rows = 100
                if len(df) > max_rows:
                    display_df = df.head(max_rows)
                    sheet_content += f"Showing first {max_rows} rows:\n"
                else:
                    display_df = df
                
                # Convert to string representation
                sheet_content += display_df.to_string(index=False, max_rows=max_rows)
                
                if len(df) > max_rows:
                    sheet_content += f"\n... and {len(df) - max_rows} more rows"
                
                # Add basic statistics for numeric columns
                numeric_columns = df.select_dtypes(include=['number']).columns
                if len(numeric_columns) > 0:
                    sheet_content += "\n\nNumeric Column Statistics:\n"
                    for col in numeric_columns:
                        sheet_content += f"{col}: Mean={df[col].mean():.2f}, Min={df[col].min()}, Max={df[col].max()}\n"
                
                all_content.append(sheet_content)
            
            if not all_content:
                raise ValueError("No data could be extracted from Excel file")
            
            return "\n".join(all_content)
        
        except Exception as e:
            raise Exception(f"Excel processing error: {str(e)}")